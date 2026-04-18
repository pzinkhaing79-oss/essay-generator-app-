import streamlit as st
import google.generativeai as genai
import docx
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders

# ---------------- UI DESIGN ---------------- #
st.set_page_config(page_title="Pro Essay Generator", layout="wide")
st.markdown("""
    <style>
    .main { background-color: #f4f6f9; color: #1e1e1e; }
    .stButton>button { background-color: #2e86c1; color: white; font-weight: bold; border-radius: 8px; width: 100%; padding: 12px; font-size: 16px; }
    h1, h2, h3 { color: #2e86c1; text-align: center; }
    </style>
    """, unsafe_allow_html=True)

st.write("### ✍️ Professional Essay & Auto-Email Generator")

# ---------------- INPUTS ---------------- #
api_key = st.text_input("🔑 Google AI Studio API Key ကို ထည့်ပါ", type="password")

st.markdown("---")
st.write("#### 📝 ခေါင်းစဉ်ရွေးချယ်ရန် (ပြိုင်ပွဲစည်းကမ်းချက်ပါ ခေါင်းစဉ် ၄ ခု)")
topics = [
    "ကျွန်ုပ်တို့၏ လူမှုပတ်ဝန်းကျင်တွင် မြေမြှုပ်မိုင်းများ ဆက်လက်အသုံးပြုမှုကို မည်သို့တားဆီးကြမည်နည်း။",
    "မိမိတို့ဒေသအတွင်း လက်နက်ကိုင်အဖွဲ့အစည်းများ၏ မိုင်းအသုံးပြုမှုကို အဆုံးသတ်နိုင်ရန် လူထုမှလုပ်ဆောင်နိုင်သည့် အဆင့်ဆင့်သော လုပ်ငန်းစဉ်များ။",
    "ကျွန်ုပ်တို့၏ လူမှုအသိုင်းအဝိုင်းကို မြေမြှုပ်မိုင်းကင်းစင်သောဇုန်အဖြစ် မည်သို့ဖော်ဆောင်မည်နည်း။",
    "မြေမြှုပ်မိုင်းကြောင့် ထိခိုက်ရသူများအတွက် ပိုမိုကောင်းမွန်သော လူမှုဘဝဖော်ဆောင်ပေးနိုင်ရန် မိမိတို့လူထုမှ မည်သို့သောလုပ်ငန်းစဉ်များ ဆောင်ရွက်နိုင်သနည်း။"
]
selected_topic = st.selectbox("ရေးသားလိုသော ခေါင်းစဉ်ကို ရွေးချယ်ပါ", topics)

st.markdown("---")
st.write("#### 📧 Auto Email ဖြင့် ပေးပို့ရန် (ရွေးချယ်နိုင်သည်)")
st.info("Email ပို့ရန် သင့် Gmail နှင့် App Password လိုအပ်ပါသည်။ (ပုံမှန် Password ဖြင့် ပို့၍မရပါ)")
sender_email = st.text_input("သင့် Gmail လိပ်စာ")
sender_password = st.text_input("Gmail App Password", type="password")
receiver_email = st.text_input("ပေးပို့မည့် လိပ်စာ", value="ArtContest@minefreemyanmar.info")

# ---------------- FUNCTIONS ---------------- #
def create_docx(text, title):
    doc = docx.Document()
    
    # Page setup (A4 Size)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    # Margins (25mm ပတ်လည်)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    
    # Title
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = heading.add_run(title)
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    
    doc.add_paragraph("") # Space
    
    # Body Text (12pt size)
    for paragraph in text.split('\n'):
        if paragraph.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(paragraph.strip())
            run.font.size = Pt(12)
            
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def send_email(sender, app_password, receiver, subject, file_stream, filename):
    msg = MIMEMultipart()
    msg['From'] = sender
    msg['To'] = receiver
    msg['Subject'] = subject
    
    body = "လေးစားအပ်ပါသော တာဝန်ရှိသူများခင်ဗျာ/ရှင့်\n\nစာစီစာကုံးပြိုင်ပွဲအတွက် ဝင်ရောက်ယှဉ်ပြိုင်မည့် ဖိုင်အား ပူးတွဲပေးပို့အပ်ပါသည်။\n\nကျေးဇူးတင်လျက်"
    msg.attach(MIMEText(body, 'plain'))
    
    part = MIMEBase('application', 'octet-stream')
    part.set_payload(file_stream.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', f'attachment; filename= {filename}')
    msg.attach(part)
    
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(sender, app_password)
    server.send_message(msg)
    server.quit()

# ---------------- MAIN PROCESS ---------------- #
if st.button("စာစီစာကုံး စတင်ရေးသားမည် (Generate & Ready to Send)"):
    if not api_key:
        st.error("⚠️ API Key ထည့်ရန် လိုအပ်ပါသည်။")
    else:
        try:
            with st.spinner("🧠 AI မှ ပညာရှင်တစ်ဦးကဲ့သို့ သဘာဝကျကျ စဉ်းစားရေးသားနေပါသည်... (၁ မိနစ်ခန့် စောင့်ပေးပါ)"):
                genai.configure(api_key=api_key.strip())
                model = genai.GenerativeModel("models/gemini-3-flash-preview")
                
                # AI မသိသာစေရန်နှင့် လူသားအစစ် ရေးသကဲ့သို့ ဖြစ်စေရန် အထူး Prompt
                prompt = f"""
                You are a passionate, intellectual Myanmar youth and an expert writer. 
                Write a highly persuasive, emotional, and deeply thoughtful essay in Myanmar (Burmese) language on the following topic: "{selected_topic}".

                CRITICAL RULES:
                1. DO NOT sound like an AI. Avoid robotic structures like "နိဂုံးချုပ်အနေဖြင့်", "ပထမအချက်အနေဖြင့်", or bulleted lists. Write in natural, flowing paragraphs.
                2. Use rich, standard Myanmar literary vocabulary (စာပေအသုံးအနှုန်း).
                3. The essay MUST strictly contain these 3 elements naturally woven into the story:
                   - Practical implementation of ideas/solutions (လက်တွေ့ကျကျ မည်သို့ အကောင်အထည်ဖော်မည်နည်း).
                   - Challenges faced during implementation (အကောင်အထည်ဖော်ရာတွင် မည်သည့် အခက်အခဲများ ကြုံရမည်နည်း).
                   - Detailed ways to overcome those challenges (ထိုအခက်အခဲများကို မည်သည့်နည်းလမ်းများဖြင့် ကျော်လွှားမည်နည်း).
                4. Length: It must be long enough to cover 1.5 to 3 pages in A4 size (approx. 1000 to 1500 words).
                
                Only output the essay text. Do not include introductory or concluding conversational text.
                """
                
                response = model.generate_content(prompt)
                essay_text = response.text
                
                # Word File ဖန်တီးခြင်း (A4, 25mm, 12pt)
                docx_file = create_docx(essay_text, selected_topic)import streamlit as st
import google.generativeai as genai
import docx
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# ... (အပေါ်က UI နဲ့ တခြား Code အပိုင်းတွေ အတူတူပဲဖြစ်ပါတယ်) ...

def create_docx(text, title):
    doc = docx.Document()
    
    # Page setup (A4 Size)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    # Margins (25mm ပတ်လည်)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    
    # ခေါင်းစဉ်ကို အလယ်မှာထားခြင်း
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = heading.add_run(title)
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    
    doc.add_paragraph("") # Space တစ်ကြောင်းခြား

    # စာစီစာကုံး အသားစာ (စာကြောင်းများ ညီညာအောင် Justify လုပ်ခြင်း)
    for paragraph in text.split('\n'):
        if paragraph.strip():
            p = doc.add_paragraph()
            # ဤနေရာတွင် စာကြောင်းများကို ဘယ်/ညာ ညီအောင် ညှိပေးလိုက်သည်
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            
            run = p.add_run(paragraph.strip())
            run.font.size = Pt(12)
            # စာပိုဒ်တစ်ခုချင်းစီရဲ့ အစမှာ Space ချန်ချင်ရင် သုံးနိုင်ပါတယ်
            p.paragraph_format.first_line_indent = Mm(12.7) 
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

                
            st.success("✅ စာစီစာကုံး ရေးသားခြင်း ပြီးမြောက်ပါပြီ!")
            
            st.write("### 📄 စာစီစာကုံး အကြမ်းဖတ်ရန်")
            st.text_area("", essay_text, height=350)
            
            st.write("### 📥 ဖိုင်ဒေါင်းလုဒ်ဆွဲရန် (Word Format - A4, 25mm, 12pt အဆင်သင့်ဖြစ်ပါသည်)")
            st.download_button(
                label="📥 စာစီစာကုံး (Word File) ကို ဒေါင်းလုဒ်ဆွဲမည်",
                data=docx_file,
                file_name="Mine_Free_Myanmar_Essay.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Email ပို့ခြင်း အပိုင်း
            if sender_email and sender_password and receiver_email:
                with st.spinner("📧 Email ပေးပို့နေပါသည်..."):
                    try:
                        docx_file.seek(0) # Reset stream for email attachment
                        send_email(sender_email, sender_password, receiver_email, f"စာစီစာကုံးပြိုင်ပွဲ ဝင်ရောက်ယှဉ်ပြိုင်ခြင်း - {selected_topic[:20]}...", docx_file, "Essay.docx")
                        st.success(f"✅ {receiver_email} သို့ အောင်မြင်စွာ ပေးပို့ပြီးပါပြီ!")
                        st.balloons()
                    except Exception as email_err:
                        st.error(f"⚠️ Email ပို့ရာတွင် အမှားဖြစ်နေပါသည်: {str(email_err)}. (Gmail App Password မှန်ကန်မှုရှိမရှိ စစ်ဆေးပါ။)")

        except Exception as e:
            st.error(f"❌ အမှားအယွင်း ဖြစ်ပေါ်နေပါသည်: {str(e)}")

